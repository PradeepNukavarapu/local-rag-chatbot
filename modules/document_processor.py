#!/usr/bin/env python3
"""
Document Processor Module
Handles document validation, text extraction, and chunking
"""

import streamlit as st
import PyPDF2
import re
from langdetect import detect, LangDetectException

# Chunking configuration
TARGET_CHUNK_SIZE = 1500
MAX_CHUNK_SIZE = 2000
OVERLAP_SIZE = 200
MIN_TEXT_LENGTH = 200

def validate_document_requirements(uploaded_file):
    """Validate document meets requirements"""
    try:
        # Check file size (max 15 MB)
        uploaded_file.seek(0, 2)
        file_size = uploaded_file.tell()
        uploaded_file.seek(0)
        
        max_size = 15 * 1024 * 1024  # 15 MB
        if file_size > max_size:
            return False, f"File size ({file_size / (1024*1024):.1f} MB) exceeds limit (15 MB)"
        
        # Check file type
        file_extension = uploaded_file.name.lower().split('.')[-1]
        if file_extension not in ['pdf', 'docx']:
            return False, f"Unsupported file type: {file_extension}. Only PDF and DOCX files are allowed."
        
        # Validate based on file type
        if file_extension == 'pdf':
            return validate_pdf_document(uploaded_file)
        elif file_extension == 'docx':
            return validate_docx_document(uploaded_file)
        
        return False, "Unsupported file type"
    except Exception as e:
        return False, f"Error validating document: {e}"

def validate_pdf_document(pdf_file):
    """Validate PDF is text-based"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_content = ""
        total_pages = len(pdf_reader.pages)
        
        # Check first 3 pages
        pages_to_check = min(3, total_pages)
        for i in range(pages_to_check):
            page = pdf_reader.pages[i]
            page_text = page.extract_text()
            
            if not page_text.strip():
                return False, f"Page {i+1} appears to be image-only. Only text-based PDFs are supported."
            
            text_content += page_text
        
        if len(text_content.strip()) < MIN_TEXT_LENGTH:
            return False, "PDF has insufficient text content."
        
        return True, f"PDF validation successful ({total_pages} pages, {len(text_content)} characters)"
    except Exception as e:
        return False, f"Error validating PDF: {e}"

def validate_docx_document(docx_file):
    """Validate DOCX contains text"""
    try:
        import docx
        doc = docx.Document(docx_file)
        text_content = ""
        
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        if len(text_content.strip()) < MIN_TEXT_LENGTH:
            return False, "DOCX has insufficient text content."
        
        return True, f"DOCX validation successful ({len(text_content)} characters)"
    except Exception as e:
        return False, f"Error validating DOCX: {e}"

def extract_text_from_document(uploaded_file):
    """Extract text from uploaded PDF or DOCX file"""
    try:
        file_extension = uploaded_file.name.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return extract_text_from_pdf(uploaded_file)
        elif file_extension == 'docx':
            return extract_text_from_docx(uploaded_file)
        else:
            st.error(f"Unsupported file type: {file_extension}")
            return None
    except Exception as e:
        st.error(f"Error extracting text from document: {e}")
        return None

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF with page tracking"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        pages_data = []
        
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text.strip():
                pages_data.append({
                    'page_number': page_num + 1,
                    'text': page_text.strip()
                })
        
        return pages_data
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return None

def extract_text_from_docx(docx_file):
    """Extract text from DOCX"""
    try:
        import docx
        doc = docx.Document(docx_file)
        pages_data = []
        
        current_section = ""
        section_count = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                current_section += text + "\n\n"
                
                # Create sections every ~1000 characters
                if len(current_section) > 1000:
                    section_count += 1
                    pages_data.append({
                        'page_number': section_count,
                        'text': current_section.strip()
                    })
                    current_section = ""
        
        # Add remaining text
        if current_section.strip():
            section_count += 1
            pages_data.append({
                'page_number': section_count,
                'text': current_section.strip()
            })
        
        return pages_data
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return None

def hybrid_chunking(text, target_size=TARGET_CHUNK_SIZE, max_size=MAX_CHUNK_SIZE, overlap=OVERLAP_SIZE):
    """Hybrid chunking: sentence-aware with size control"""
    if not text or len(text) <= target_size:
        return [text] if text else []
    
    chunks = []
    current_chunk = ""
    
    # Split by paragraphs first
    paragraphs = text.split('\n\n')
    
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        
        # If adding this paragraph exceeds max size, save current chunk
        if len(current_chunk) + len(paragraph) > max_size and current_chunk:
            chunks.append(current_chunk.strip())
            # Start new chunk with overlap
            overlap_text = current_chunk[-overlap:] if overlap > 0 else ""
            current_chunk = overlap_text + "\n\n" + paragraph
        else:
            current_chunk += "\n\n" + paragraph if current_chunk else paragraph
    
    # Add the last chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    # Final safety check - split very large chunks by sentences
    final_chunks = []
    for chunk in chunks:
        if len(chunk) <= max_size:
            final_chunks.append(chunk)
        else:
            # Split by sentences
            sentences = re.split(r'(?<=[.!?])\s+', chunk)
            temp_chunk = ""
            
            for sentence in sentences:
                if len(temp_chunk) + len(sentence) > max_size and temp_chunk:
                    final_chunks.append(temp_chunk.strip())
                    temp_chunk = sentence
                else:
                    temp_chunk += " " + sentence if temp_chunk else sentence
            
            if temp_chunk.strip():
                final_chunks.append(temp_chunk.strip())
    
    return final_chunks